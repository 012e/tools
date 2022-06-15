import re
from string import ascii_uppercase
from docx import Document
from docx.shared import RGBColor

question_regex = re.compile(r'<[^"]*-COA>', re.X)


def is_good_question(question: str) -> bool:
    return not "VDT" in question


question_content_regex = re.compile(r'(?<=>\ ).*$', re.X)

all_answer_regex = re.compile(r'(<.>)', re.X)

ANS_CHAR = dict(zip(range(0, 26), ascii_uppercase))


def is_end(text):
    return text == "<END>"


def add_bold(line, content):
    state = line.add_run(content)
    state.bold = True


def add_red_bold(line, content):
    state = line.add_run(content)
    state.bold = True
    state.font.color.rgb = RGBColor(255, 0, 0)


def add_normal(line, content):
    line.add_run(content)


def key_lookup(key: int) -> str:
    return ANS_CHAR[key]


def next_key() -> str:
    global ans_count, current_key
    ans_count += 1
    if ans_count == 4:
        ans_count = 0

    current_key = key_lookup(ans_count) + '. '


def add_correct_answer(answer: str):
    global line, current_key
    key = current_key
    add_red_bold(line, f"{key} {answer}")


def add_normal_answer(answer: str):
    global line, current_key
    key = current_key
    add_normal(line, f"{key} {answer}")


def handle_answer(text: str):
    global current_key
    all = re.split(all_answer_regex, text)
    if "" in all:
        all.remove("")

    token_is_corrent_answer = True
    for index, token in enumerate(all):
        match token:
            case "<#>":
                token_is_corrent_answer = True
                next_key()
                continue
            case "<$>":
                token_is_corrent_answer = False
                next_key()
                continue
            case "\t":
                continue
        if token_is_corrent_answer:
            add_correct_answer(token)
        else:
            add_normal_answer(token)


def is_answer(text: str) -> bool:
    return "<#>" in text or "<$>" in text


def is_question(text):
    return re.search(question_regex, text)


def handle_question(text: str):
    global question_count
    if is_question(text):
        if is_good_question(text):
            add_bold(line, f"Câu {question_count}. ")
        else:
            add_bold(line, f"Câu {question_count} (VDT). ")
        add_normal(line, question_content(text))
        question_count += 1
    else:
        add_normal(line, text)


def question_content(text: str) -> str:
    return re.findall(question_content_regex, text)


def handle_end():
    line.text = ""


def delete_line():
    p = line._element
    p.getparent().remove(p)
    p._p = p._element = None


ALL_FILES = ["ancolphenol", "onthi", "75lan4", "75lan3", "75lan2"]
for file_name in ALL_FILES:
    document = Document('./' + file_name + ".docx")

    ans_count = -1
    question_count = 1
    current_key = 'A. '

    for i, line in enumerate(document.paragraphs):
        text = line.text
        line.text = ""  # reset line

        if text == "":
            delete_line()
            continue

        if is_end(text):
            delete_line()
            continue

        if is_answer(text):
            handle_answer(text)
            continue

        handle_question(text)

    document.save(file_name + '-converted.docx')
